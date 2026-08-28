from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parent / "Relatorio_Gerencial_SalesTrack_2026-08-28.docx"

NAVY = "182235"
BLUE = "2F6FED"
TEAL = "08A88A"
GOLD = "F2B705"
RED = "D6455D"
GRAY = "5C667A"
LIGHT = "F3F6FA"
PALE_BLUE = "EAF1FF"
PALE_GREEN = "E8F7F2"
PALE_GOLD = "FFF6D9"
PALE_RED = "FDECEF"
WHITE = "FFFFFF"
BLACK = "17202A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Arial", size=10.5, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in (("Heading 1", 16, 14, 7), ("Heading 2", 13, 10, 5), ("Heading 3", 11, 8, 4)):
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE if name != "Heading 3" else NAVY)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.base_style = styles["Normal"]
        callout.font.name = "Arial"
        callout.font.size = Pt(11)
        callout.font.bold = True
        callout.font.color.rgb = RGBColor.from_string(NAVY)
        callout.paragraph_format.left_indent = Inches(0.2)
        callout.paragraph_format.right_indent = Inches(0.2)
        callout.paragraph_format.space_before = Pt(6)
        callout.paragraph_format.space_after = Pt(10)


def setup_header_footer(section):
    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("SALESTRACK  |  RELATÓRIO GERENCIAL")
        set_font(r, size=8.5, color=GRAY, bold=True)

    for footer in (section.footer, section.even_page_footer):
        table = footer.add_table(rows=1, cols=2, width=Inches(6.4))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(table, [5.25, 1.15])
        left = table.cell(0, 0).paragraphs[0]
        r = left.add_run("Uso interno • Projeto SalesTrack • 28/08/2026")
        set_font(r, size=8, color=GRAY)
        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = right.add_run("Página ")
        set_font(r, size=8, color=GRAY)
        add_field(right, "PAGE")


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("RELATÓRIO GERENCIAL")
    set_font(r, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Projeto SalesTrack")
    set_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Dashboard comercial, inteligência de vendas e governança operacional")
    set_font(r, size=14, color=GRAY)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_col_widths(meta, [1.45, 5.25])
    rows = [
        ("Data de referência", "28 de agosto de 2026"),
        ("Status", "Versão operacional validada em navegador"),
        ("Fonte principal", "Google Sheets publicado em CSV"),
        ("Objetivo", "Apoiar decisões comerciais, operacionais e de comissionamento"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], PALE_BLUE)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_font(p1.add_run(label), size=9, color=NAVY, bold=True)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run(value), size=9.5, color=BLACK)

    doc.add_paragraph()
    p = doc.add_paragraph(style="Callout")
    r = p.add_run("Síntese executiva: ")
    set_font(r, size=11, color=BLUE, bold=True)
    r = p.add_run("o projeto evoluiu de uma leitura básica de planilha para um painel gerencial responsivo, com regras próprias de negócio, filtros comparativos, visão de volumes, clientes, produtos, vendedores, comissões e cancelamentos.")
    set_font(r, size=11, color=NAVY, bold=True)
    set_cell_like_paragraph(p, PALE_BLUE, BLUE)


def set_cell_like_paragraph(p, fill, border):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_priority_table(doc):
    doc.add_page_break()
    doc.add_heading("2. Avaliação por nível de impacto", level=1)
    p = doc.add_paragraph("A classificação abaixo considera impacto gerencial, aderência ao negócio e risco residual da solução atual.")
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_col_widths(table, [1.0, 1.65, 2.45, 1.6])
    headers = ["Nível", "Tema", "Leitura gerencial", "Diretriz"]
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_cell_margins(table.rows[0].cells[i])
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("ALTO", "Valor entregue", "Visão única do período, comparação temporal, volumes em bandejas/kg/toneladas, comissões, vendedores, clientes e produtos.", "Manter como núcleo decisório."),
        ("ALTO", "Aderência ao negócio", "Regras específicas foram incorporadas: coluna I como valor líquido, venda única pela coluna C, semana domingo–sábado e conversões independentes.", "Documentar e testar regressões."),
        ("ALTO", "Usabilidade executiva", "Filtros ERP-like, KPIs comparativos, indicadores de crescimento/queda e grade 16:9 alinhada.", "Preservar simplicidade e desempenho."),
        ("MÉDIO", "Governança de dados", "A fonte é CSV publicado; mudanças de cabeçalho, status ou estrutura podem afetar a interpretação.", "Criar validação de esquema e aviso de qualidade."),
        ("MÉDIO", "Consolidação cadastral", "Matrizes e filiais são agrupadas por regras textuais para Zona Sul, Hortifruti, Rede Ultra, Cardin, Temakeria, Domenica e Aipo e Aipim.", "Centralizar cadastro mestre."),
        ("MÉDIO", "Portabilidade", "O painel funciona via servidor local e consulta a planilha publicada, mas depende de internet, publicação ativa e navegador compatível.", "Empacotar instalação e manual de contingência."),
        ("BAIXO", "Escala atual", "O volume observado é compatível com processamento local, desde que gráficos permaneçam limitados e instâncias anteriores sejam destruídas.", "Monitorar memória e tempo de carga."),
        ("BAIXO", "Manutenção visual", "A grade compartilhada reduz desalinhamentos; ajustes futuros ainda exigem teste em resoluções distintas.", "Criar checklist visual responsivo."),
    ]
    fills = {"ALTO": PALE_GREEN, "MÉDIO": PALE_GOLD, "BAIXO": LIGHT}
    colors = {"ALTO": TEAL, "MÉDIO": "9A6A00", "BAIXO": GRAY}
    for level, theme, reading, direction in rows:
        cells = table.add_row().cells
        vals = [level, theme, reading, direction]
        for idx, val in enumerate(vals):
            set_cell_margins(cells[idx], top=80, bottom=80)
            if idx == 0:
                set_cell_shading(cells[idx], fills[level])
            set_font(cells[idx].paragraphs[0].add_run(val), size=8.6, color=colors[level] if idx == 0 else BLACK, bold=(idx in (0, 1)))


def add_swot(doc):
    doc.add_page_break()
    doc.add_heading("3. Matriz SWOT", level=1)
    p = doc.add_paragraph("A matriz combina capacidades já entregues com fatores externos e riscos de continuidade.")
    p.paragraph_format.space_after = Pt(9)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(table, [3.35, 3.35])
    blocks = [
        (0, 0, "FORÇAS", PALE_GREEN, TEAL, [
            "Regras de negócio incorporadas diretamente ao painel.",
            "Leitura do valor líquido pela coluna I e vendas únicas pela coluna C.",
            "Exclusão de cancelamentos da visão principal, com modo analítico específico.",
            "Comparações temporais coerentes com filtros de dia, semana, mês, ano e períodos móveis.",
            "Design responsivo, legível e orientado a decisão." ]),
        (0, 1, "FRAQUEZAS", PALE_RED, RED, [
            "Dependência de CSV público e conectividade.",
            "Regras cadastrais de clientes ainda mantidas no código.",
            "Ausência de autenticação, perfis e trilha formal de auditoria.",
            "Processamento integral no navegador limita escala futura.",
            "Testes automatizados ainda não institucionalizados." ]),
        (1, 0, "OPORTUNIDADES", PALE_BLUE, BLUE, [
            "Transformar o painel em portal comercial multiusuário.",
            "Criar metas, previsão, margem, mix e alertas de queda.",
            "Integrar ERP/API e eliminar dependência de publicação manual.",
            "Adicionar cadastro mestre para clientes, produtos e vendedores.",
            "Distribuir relatórios executivos periódicos." ]),
        (1, 1, "AMEAÇAS", PALE_GOLD, "9A6A00", [
            "Mudança silenciosa de colunas, nomes ou status na origem.",
            "Indisponibilidade do Google Sheets ou revogação do link.",
            "Exposição indevida de dados por publicação aberta.",
            "Decisões erradas se conversões ou comissões forem alteradas sem governança.",
            "Degradação do navegador com crescimento expressivo da base." ]),
    ]
    for r_idx, c_idx, title, fill, color, bullets in blocks:
        cell = table.cell(r_idx, c_idx)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=140, start=150, bottom=140, end=150)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(title), size=11, color=color, bold=True)
        for item in bullets:
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            set_font(p.add_run(item), size=8.6, color=BLACK)


def add_crisis(doc):
    doc.add_page_break()
    doc.add_heading("4. Gestão de crise e continuidade", level=1)
    p = doc.add_paragraph(style="Callout")
    set_font(p.add_run("Princípio: "), size=11, color=RED, bold=True)
    set_font(p.add_run("em uma divergência de dados, a prioridade não é manter o dashboard no ar a qualquer custo; é impedir que uma informação não validada gere decisão comercial, financeira ou de pagamento."), size=11, color=NAVY, bold=True)
    set_cell_like_paragraph(p, PALE_RED, RED)

    doc.add_heading("Gatilhos de crise", level=2)
    for text in [
        "Dashboard sem atualização, CSV inacessível ou contagem de registros anormal.",
        "Diferença entre o valor líquido do painel e a conferência da coluna I.",
        "Venda cancelada aparecendo em faturamento, volumes ou comissões.",
        "Comissão divergente, vendedor atribuído ao período incorreto ou venda duplicada.",
        "Uso excessivo de memória, travamento do Chrome ou gráfico crescendo indefinidamente.",
        "Suspeita de exposição indevida do link público ou dos dados comerciais.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("Protocolo de resposta", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(table, [0.65, 1.35, 3.35, 1.35])
    for i, h in enumerate(["Fase", "Prazo", "Ação", "Responsável"]):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_cell_margins(table.rows[0].cells[i])
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(h), size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    phases = [
        ("1", "0–15 min", "Congelar decisões e pagamentos baseados no painel; registrar filtro, horário, print e sintoma.", "Usuário-chave"),
        ("2", "15–30 min", "Validar disponibilidade da planilha, cabeçalhos, data final inclusiva, status e amostra da coluna I.", "Dono dos dados"),
        ("3", "30–60 min", "Comparar dashboard com fonte em amostra controlada; classificar impacto financeiro, operacional e de privacidade.", "Gestor + suporte"),
        ("4", "Até 2 h", "Restaurar último backup estável ou aplicar correção isolada; testar período afetado e período anterior.", "Responsável técnico"),
        ("5", "Mesmo dia", "Comunicar liberação, escopo afetado, período de risco e orientação de reconciliação.", "Gestor do sistema"),
        ("6", "Até 48 h", "Emitir causa raiz, controles preventivos, responsável e prazo de conclusão.", "Comitê responsável"),
    ]
    for phase in phases:
        cells = table.add_row().cells
        for idx, val in enumerate(phase):
            set_cell_margins(cells[idx], top=80, bottom=80)
            set_font(cells[idx].paragraphs[0].add_run(val), size=8.6, bold=(idx == 0))

    doc.add_heading("Regras de comunicação", level=2)
    for text in [
        "Usar uma única mensagem oficial, informando fato, impacto, ação e próxima atualização.",
        "Evitar comunicar números não reconciliados ou atribuir causa antes da validação.",
        "Manter registro dos filtros, versões, backups e evidências usados na correção.",
        "Quando houver impacto em comissão, suspender o pagamento afetado até dupla conferência.",
    ]:
        add_bullet(doc, text)


def add_roadmap(doc):
    doc.add_page_break()
    doc.add_heading("5. Recomendações e plano 30–60–90 dias", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(table, [0.9, 2.75, 1.55, 1.5])
    for i, h in enumerate(["Horizonte", "Entregas recomendadas", "Resultado esperado", "Prioridade"]):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_cell_margins(table.rows[0].cells[i])
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(h), size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("0–30 dias", "Formalizar dicionário de dados; criar testes para colunas C, G e I; registrar regras de cancelamento, conversão e comissão; documentar restauração de backup.", "Confiabilidade e rastreabilidade.", "ALTA"),
        ("31–60 dias", "Externalizar cadastro mestre de clientes/produtos; criar validação de esquema; exibir data/hora da última atualização e alerta de defasagem.", "Menor risco de manutenção manual.", "ALTA"),
        ("61–90 dias", "Avaliar API/ERP, autenticação, banco histórico, exportação gerencial e monitoramento de desempenho.", "Escalabilidade e segurança.", "MÉDIA"),
    ]
    for horizon, deliver, result, priority in rows:
        cells = table.add_row().cells
        for idx, val in enumerate((horizon, deliver, result, priority)):
            set_cell_margins(cells[idx], top=100, bottom=100)
            set_font(cells[idx].paragraphs[0].add_run(val), size=8.8, bold=(idx in (0, 3)), color=RED if priority == "ALTA" and idx == 3 else BLACK)

    doc.add_heading("Indicadores de governança sugeridos", level=2)
    for text in [
        "Taxa de atualização bem-sucedida e tempo desde a última carga válida.",
        "Diferença entre total do painel e amostra reconciliada da coluna I.",
        "Número de vendas únicas sem identificador na coluna C.",
        "Quantidade de linhas com unidade ou produto não reconhecido.",
        "Tempo médio de recuperação e número de incidentes por versão.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("6. Conclusão gerencial", level=1)
    p = doc.add_paragraph()
    set_font(p.add_run("O SalesTrack já atingiu valor operacional relevante: "), bold=True, color=NAVY)
    set_font(p.add_run("reduz dispersão de informações, traduz regras específicas da empresa e melhora a leitura de desempenho por período. O próximo salto não depende principalmente de novos gráficos, mas de governança da origem, testes automatizados, segurança de acesso e formalização da continuidade operacional."))
    p = doc.add_paragraph(style="Callout")
    set_font(p.add_run("Recomendação executiva: "), size=11, color=TEAL, bold=True)
    set_font(p.add_run("manter o painel em operação controlada, aprovar a etapa de governança de dados e instituir o protocolo de crise antes de ampliar o uso para decisões financeiras recorrentes."), size=11, color=NAVY, bold=True)
    set_cell_like_paragraph(p, PALE_GREEN, TEAL)


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    style_doc(doc)
    setup_header_footer(doc.sections[0])
    add_title_block(doc)

    doc.add_page_break()
    doc.add_heading("1. Escopo e principais entregas", level=1)
    intro = doc.add_paragraph("O desenvolvimento realizado consolidou regras de negócio, experiência visual e controles operacionais em um único dashboard. As entregas mais relevantes foram:")
    intro.paragraph_format.space_after = Pt(6)
    deliverables = [
        "Integração com CSV publicado pelo Google Sheets e operação via servidor local no Chrome.",
        "Tratamento de codificação, datas brasileiras e inclusão correta da data final do filtro.",
        "Valor líquido estritamente baseado na coluna I e contagem de vendas únicas pela coluna C.",
        "Exclusão de cancelamentos da visão principal, com modo interativo dedicado à análise de cancelamentos.",
        "Conversões internas de bandejas e quilogramas, incluindo 205 g por bandeja e separação de Shiitake inteiro/fatiado.",
        "KPIs de vendas, toneladas, bandejas, quilos de Shiitake, cancelamentos e melhor vendedor, todos com comparação temporal.",
        "Resumo por vendedor com número de vendas, total vendido, média por venda, comissão e variações frente ao período anterior.",
        "Consolidação de filiais por matriz na participação de clientes e legendas com valor e percentual.",
        "Filtros de período inspirados no ERP, incluindo semana de domingo a sábado e navegação contextual por setas.",
        "Gráficos com limites de altura, barras comparativas e grade 16:9 compartilhada para reduzir uso de memória e manter alinhamento.",
        "Rotina disciplinada de backups antes de cada substituição relevante.",
    ]
    for item in deliverables:
        add_bullet(doc, item)

    doc.add_heading("Leitura de maturidade", level=2)
    p = doc.add_paragraph()
    set_font(p.add_run("Situação atual: "), bold=True, color=BLUE)
    set_font(p.add_run("MVP operacional avançado / ferramenta gerencial interna."))
    p = doc.add_paragraph()
    set_font(p.add_run("Limite atual: "), bold=True, color=RED)
    set_font(p.add_run("a solução ainda não deve ser tratada como sistema financeiro auditado ou plataforma corporativa com controle de acesso."))

    add_priority_table(doc)
    add_swot(doc)
    add_crisis(doc)
    add_roadmap(doc)

    props = doc.core_properties
    props.title = "Relatório Gerencial — Projeto SalesTrack"
    props.subject = "Avaliação executiva, SWOT e gestão de crise"
    props.author = "SalesTrack"
    props.keywords = "SalesTrack, dashboard, vendas, gestão, SWOT, crise"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
